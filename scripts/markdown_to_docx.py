#!/usr/bin/env python3
"""
DOCX Report Generator
Converts markdown report to formatted .docx document.
"""

import sys
import json
import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class DocxReportGenerator:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configure document styles."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.space_before = Pt(0)

        # Heading styles
        for level in range(1, 4):
            heading_style = self.doc.styles[f'Heading {level}']
            heading_style.font.name = 'Calibri'
            heading_style.font.color.rgb = RGBColor(0, 51, 102)
            if level == 1:
                heading_style.font.size = Pt(20)
                heading_style.font.bold = True
            elif level == 2:
                heading_style.font.size = Pt(16)
                heading_style.font.bold = True
            elif level == 3:
                heading_style.font.size = Pt(13)
                heading_style.font.bold = True

    def add_title_page(self, company_name: str, ticker: str, filing_date: str, report_date: str, cik: str):
        """Add a professional title page."""
        # Add spacing
        for _ in range(4):
            self.doc.add_paragraph()

        # Company name
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(company_name)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        # Ticker
        ticker_para = self.doc.add_paragraph()
        ticker_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ticker_para.add_run(f"({ticker})")
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(100, 100, 100)

        # Subtitle
        self.doc.add_paragraph()
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("SEC Form 10-K Analysis Report")
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(80, 80, 80)

        # Details
        self.doc.add_paragraph()
        details = self.doc.add_paragraph()
        details.alignment = WD_ALIGN_PARAGRAPH.CENTER
        details_text = (
            f"Filing Date: {filing_date}\n"
            f"Period End: {report_date}\n"
            f"CIK: {cik}\n"
            f"Report Generated: {datetime.now().strftime('%B %d, %Y')}"
        )
        run = details.add_run(details_text)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(100, 100, 100)

        # Page break
        self.doc.add_page_break()

    def add_table_of_contents(self):
        """Add a table of contents placeholder."""
        self.doc.add_heading('Table of Contents', level=1)
        p = self.doc.add_paragraph()
        run = p.add_run('Right-click and select "Update Field" to refresh the table of contents.')
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        self.doc.add_page_break()

    def parse_markdown_to_docx(self, markdown_text: str):
        """Convert markdown text to docx elements."""
        lines = markdown_text.split('\n')
        i = 0
        in_table = False
        table_rows = []
        table_headers = None

        while i < len(lines):
            line = lines[i]

            # Headings
            if line.startswith('# '):
                self._flush_table(table_headers, table_rows)
                self.doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                self._flush_table(table_headers, table_rows)
                self.doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                self._flush_table(table_headers, table_rows)
                self.doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                self._flush_table(table_headers, table_rows)
                self.doc.add_heading(line[5:], level=4)

            # Horizontal rule
            elif line.strip() == '---':
                self._flush_table(table_headers, table_rows)
                self.doc.add_paragraph().add_run().add_break()

            # Tables
            elif line.startswith('|') and line.endswith('|'):
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if not in_table:
                    in_table = True
                    table_headers = cells
                    table_rows = []
                else:
                    # Check if this is a separator row
                    if all(re.match(r'^:?-+:?$', cell) for cell in cells):
                        pass  # Skip separator row
                    else:
                        table_rows.append(cells)
            else:
                self._flush_table(table_headers, table_rows)
                in_table = False
                table_headers = None
                table_rows = []

                # Bold/italic text
                if line.strip():
                    p = self.doc.add_paragraph()
                    self._add_formatted_text(p, line)

            i += 1

        self._flush_table(table_headers, table_rows)

    def _flush_table(self, headers, rows):
        """Create a table from accumulated rows."""
        if headers and rows:
            table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header row
            for j, header in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = header
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)

            # Data rows
            for i, row_data in enumerate(rows):
                for j, cell_text in enumerate(row_data):
                    if j < len(table.rows[i + 1].cells):
                        cell = table.rows[i + 1].cells[j]
                        cell.text = cell_text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)

            self.doc.add_paragraph()  # Space after table

    def _add_formatted_text(self, paragraph, text):
        """Add text with markdown formatting (bold, italic)."""
        # Simple regex for **bold** and *italic*
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                run = paragraph.add_run(part)

    def save(self, filepath: str):
        """Save the document."""
        self.doc.save(filepath)


def markdown_to_docx(markdown_file: str, docx_file: str, company_name: str, ticker: str,
                     filing_date: str, report_date: str, cik: str):
    """Convert markdown file to docx."""
    # Try different encodings
    encodings = ['utf-8', 'utf-8-sig', 'cp1252', 'latin-1']
    markdown_content = None

    for encoding in encodings:
        try:
            with open(markdown_file, 'r', encoding=encoding) as f:
                markdown_content = f.read()
            break
        except UnicodeDecodeError:
            continue

    if markdown_content is None:
        raise ValueError(f"Could not decode {markdown_file} with any supported encoding")

    generator = DocxReportGenerator()
    generator.add_title_page(company_name, ticker, filing_date, report_date, cik)
    generator.add_table_of_contents()
    generator.parse_markdown_to_docx(markdown_content)
    generator.save(docx_file)

    print(f"DOCX report saved to {docx_file}", file=sys.stderr)


def main():
    if len(sys.argv) < 7:
        print("Usage: python markdown_to_docx.py <markdown_file> <docx_file> <company_name> <ticker> <filing_date> <report_date> <cik>", file=sys.stderr)
        sys.exit(1)

    markdown_file = sys.argv[1]
    docx_file = sys.argv[2]
    company_name = sys.argv[3]
    ticker = sys.argv[4]
    filing_date = sys.argv[5]
    report_date = sys.argv[6]
    cik = sys.argv[7] if len(sys.argv) > 7 else ""

    try:
        markdown_to_docx(markdown_file, docx_file, company_name, ticker, filing_date, report_date, cik)
        print(json.dumps({"success": True, "output_file": docx_file}))
    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)


if __name__ == "__main__":
    main()